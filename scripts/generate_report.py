#!/usr/bin/env python3
"""Generate the final SLR report as a Word document."""
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = Path(r"C:\Users\cpanj\OneDrive\Desktop\DTE\AgenticAI\slr-agent")

# Load data
with open(BASE / "data/05_extracted/extraction.jsonl") as f:
    extractions = [json.loads(l) for l in f if l.strip()]

with open(BASE / "data/05_extracted/quality.jsonl") as f:
    quality = {q["id"]: q for q in (json.loads(l) for l in f if l.strip())}

synthesis_text = (BASE / "output/synthesis.md").read_text(encoding="utf-8")
protocol_text = (BASE / "protocol.md").read_text(encoding="utf-8")

# Filter out insufficient
papers = [e for e in extractions if e.get("extraction_status") != "insufficient_text"]

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level, size in [(1, 16), (2, 14), (3, 12)]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Arial'
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# ===================== TITLE PAGE =====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(120)
run = p.add_run("Systematic Literature Review")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("What Effect Do Platform Strategies Have\nin Biotech Companies?")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(40)
run = p.add_run("PRISMA 2020 Compliant Report")
run.font.size = Pt(12)
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(20)
run = p.add_run("June 2026")
run.font.size = Pt(12)

doc.add_page_break()

# ===================== STRUCTURED ABSTRACT =====================
doc.add_heading("Abstract", level=1)

abstract_sections = {
    "Background": "Platform strategies are increasingly adopted by biotechnology firms, yet no systematic review has synthesized the empirical and conceptual evidence on their effects. Understanding whether and under what conditions platform strategies create value for biotech companies is critical for both managers and policymakers.",
    "Objective": "To systematically review the literature on the effects of platform strategies in biotechnology companies, addressing four sub-questions: (1) what types of platform strategies are observed, (2) what firm-level outcomes are associated, (3) under what conditions do they create or destroy value, and (4) what boundary conditions and trade-offs are documented.",
    "Methods": "Following PRISMA 2020 guidelines, we searched OpenAlex, Semantic Scholar, EuropePMC, and CrossRef for English-language peer-reviewed studies published from 2005 onward. After deduplication (3,879 records), keyword pre-filtering (3,793), and title/abstract screening, 92 papers were identified for inclusion. Full text was retrieved for 53; 52 underwent data extraction and MMAT quality assessment. Thematic synthesis followed Thomas and Harden (2008).",
    "Results": "The corpus is predominantly qualitative (48/52 papers). Technology platforms and R&D/open innovation platforms are the most observed types; multi-sided and modular platforms are largely absent. Effects are overwhelmingly conditional (12 papers) rather than universally positive (5) or negative (2). Key moderators include ecosystem maturity, governance design, regulatory environment, and firm dynamic capabilities. No study demonstrates a robust positive causal effect of platform strategy on biotech firm financial performance.",
    "Conclusions": "Platform strategies in biotech produce conditional rather than universal benefits. The evidence base is thin, methodologically weak, and dominated by qualitative case studies. Ten specific research gaps are identified, including the absence of platform-vs-non-platform financial comparisons and the neglect of modular product platforms."
}

for heading, text in abstract_sections.items():
    p = doc.add_paragraph()
    run = p.add_run(f"{heading}: ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)

doc.add_page_break()

# ===================== 1. BACKGROUND =====================
doc.add_heading("1. Background", level=1)

doc.add_paragraph(
    "Platform strategies have transformed competitive dynamics across industries, from technology to financial services. In biotechnology, firms increasingly organize around technology platforms, open innovation ecosystems, and data-sharing infrastructure to accelerate drug discovery, reduce development costs, and form strategic alliances. Yet the theoretical and empirical literature on platform strategies has largely developed in the context of digital platforms and consumer-facing markets, with limited attention to the distinctive features of the biotech sector: long development timelines, regulatory complexity, high capital intensity, and deep scientific uncertainty."
)
doc.add_paragraph(
    "No prior systematic review has synthesized the evidence on how platform strategies affect biotech firms. This review addresses that gap by asking: What effect do platform strategies have in biotech companies? We decompose this into four sub-questions covering platform types, firm-level outcomes, value-creation conditions, and documented boundary conditions and trade-offs."
)

# ===================== 2. METHODS =====================
doc.add_heading("2. Methods", level=1)

doc.add_heading("2.1 Protocol and Registration", level=2)
doc.add_paragraph(
    "This review follows the PRISMA 2020 statement. The protocol was developed prior to the search and specifies the research question, PICOC framing, inclusion/exclusion criteria, databases, search strategy, data extraction schema, quality assessment tool, and synthesis approach."
)

doc.add_heading("2.2 PICOC Framing", level=2)
picoc = [
    ("Population", "Biotechnology firms (therapeutics, diagnostics, agbiotech, industrial biotech). Pure pharma incumbents excluded unless analysed as a comparison group."),
    ("Intervention", "Adoption or use of a platform strategy (technology platform, multi-sided platform, modular product platform, R&D platform)."),
    ("Comparison", "Non-platform / pipeline / single-asset strategies, or pre/post adoption."),
    ("Outcome", "Innovation output, financial performance, valuation, partnership/alliance formation, survival, time-to-market, ecosystem dynamics."),
    ("Context", "Any geography; commercial biotech firms (publicly or privately held).")
]
for label, desc in picoc:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("2.3 Databases and Search Strategy", level=2)
doc.add_paragraph(
    "Four databases were searched: OpenAlex (primary, broadest recall), Semantic Scholar (cross-validation), EuropePMC (biomedical depth and OA full text), and CrossRef (metadata reconciliation and DOI canonicalization). The search combined three concept clusters using Boolean logic:"
)
doc.add_paragraph(
    'Cluster 1 (Platform): "platform strategy" OR "platform business model" OR "platform ecosystem" OR "innovation platform" OR "technology platform" OR "digital platform" OR "platform-based" OR "modular platform"',
    style='List Bullet'
)
doc.add_paragraph(
    'Cluster 2 (Biotech): "biotechnology" OR "biotech" OR "biopharmaceutical" OR "biopharma" OR "life sciences"',
    style='List Bullet'
)
doc.add_paragraph(
    'Cluster 3 (Outcomes): "firm performance" OR "competitive advantage" OR "value creation" OR "strategic alliance" OR "drug development" OR "open innovation"',
    style='List Bullet'
)
doc.add_paragraph("Date filter: 2005-01-01 to present. Language: English.")

doc.add_heading("2.4 Inclusion and Exclusion Criteria", level=2)
doc.add_paragraph("Inclusion criteria (all must hold):")
for ic in [
    "Peer-reviewed journal article, conference paper, or established working paper",
    "Empirical (qualitative, quantitative, mixed) or substantive conceptual/theory paper",
    "Explicitly addresses platform strategy AND a biotech/life sciences context",
    "Published 2005 or later",
    "English language",
    "Full text obtainable"
]:
    doc.add_paragraph(ic, style='List Bullet')

doc.add_paragraph("Exclusion criteria (any triggers exclusion):")
for ec in [
    "Editorial, book review, news item, blog post",
    'Platform strategy discussed only as background; biotech case is only illustrative',
    '"Platform" used purely in the technical/laboratory sense without strategic content',
    "Pharma-only with no biotech firms in the sample",
    "Pre-2005 or non-English",
    "Full text inaccessible after two retrieval attempts"
]:
    doc.add_paragraph(ec, style='List Bullet')

doc.add_heading("2.5 Screening Procedure", level=2)
doc.add_paragraph(
    "After deduplication, records were pre-filtered using keyword heuristics (high-recall filter dropping only records with no platform or biotech terms and under 100 characters of text). Title/abstract screening applied inclusion/exclusion criteria using a two-pass approach: an initial heuristic screen followed by a second-pass screen on MAYBE records using richer abstract-text analysis. Records with no available abstract were excluded in the second pass."
)

doc.add_heading("2.6 Data Extraction", level=2)
doc.add_paragraph(
    "Data were extracted from full-text papers using keyword-based heuristics covering: bibliographic details, theoretical lens, method, sample characteristics, platform definition, platform type, outcomes measured, effect direction and summary, boundary conditions, and stated research gaps."
)

doc.add_heading("2.7 Quality Assessment", level=2)
doc.add_paragraph(
    "The Mixed Methods Appraisal Tool (MMAT) 2018 was applied to each paper, scoring five criteria (0-1 each, total 0-5) appropriate to the method category. Quality scores are reported alongside findings but do not gate inclusion."
)

doc.add_heading("2.8 Synthesis Approach", level=2)
doc.add_paragraph(
    "Thematic synthesis following Thomas and Harden (2008): line-by-line coding of findings, grouping into descriptive themes, then mapping to analytical themes aligned with the four sub-questions."
)

doc.add_page_break()

# ===================== 3. RESULTS =====================
doc.add_heading("3. Results", level=1)

doc.add_heading("3.1 PRISMA Flow", level=2)
doc.add_paragraph(
    "Figure 1 presents the PRISMA flow diagram (see output/prisma_flow.svg). A total of 4,008 records were identified across databases (OpenAlex: 2,000; EuropePMC: 8; CrossRef: 2,000; Semantic Scholar: 0). After deduplication (3,879 unique), keyword pre-filtering (3,793), and two-pass screening, 92 papers were included. Full text was retrieved for 53 papers; 52 underwent extraction and quality assessment."
)

# PRISMA flow table
flow_data = [
    ("Records identified", "4,008"),
    ("After deduplication", "3,879"),
    ("After keyword pre-filter", "3,793"),
    ("After title/abstract screening (INCLUDE)", "92"),
    ("Full text retrieved", "53"),
    ("Data extracted", "52"),
    ("Included in synthesis", "52"),
]
table = doc.add_table(rows=len(flow_data) + 1, cols=2)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Stage"
hdr[1].text = "N"
for i, (stage, n) in enumerate(flow_data):
    table.rows[i + 1].cells[0].text = stage
    table.rows[i + 1].cells[1].text = n

doc.add_paragraph()

doc.add_heading("3.2 Corpus Descriptive Characteristics", level=2)
doc.add_paragraph(
    "The corpus comprises 52 papers. The majority cluster post-2020 among records with parseable years, though 34 records lack explicit publication years. Geographically, the USA (4), UK (5), and Europe broadly (4) dominate, with most records lacking geography data. The corpus is overwhelmingly qualitative: 48 papers use qualitative or conceptual methods, 3 use quantitative descriptive approaches, and 1 uses a quantitative non-randomized design (event study)."
)

# Descriptive characteristics table
desc_data = [
    ("Method", ""),
    ("  Qualitative/case study", "48"),
    ("  Quantitative descriptive", "3"),
    ("  Quantitative non-randomized", "1"),
    ("Theoretical lens", ""),
    ("  Ecosystem theory", "13"),
    ("  Dynamic capabilities", "4"),
    ("  Platform theory", "3"),
    ("  Resource-based view", "1"),
    ("  Not reported", "31"),
    ("Platform type", ""),
    ("  Technology platform", "6"),
    ("  Multi-sided platform", "1"),
    ("  Other/not classified", "45"),
    ("Effect direction", ""),
    ("  Conditional", "12"),
    ("  Positive", "5"),
    ("  Mixed", "4"),
    ("  Negative", "2"),
    ("  Not coded", "29"),
]
table = doc.add_table(rows=len(desc_data) + 1, cols=2)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Characteristic"
hdr[1].text = "Count"
for i, (char, cnt) in enumerate(desc_data):
    table.rows[i + 1].cells[0].text = char
    table.rows[i + 1].cells[1].text = cnt

doc.add_paragraph()

# ===================== 3.3 Findings by Sub-question =====================
doc.add_heading("3.3 Findings by Sub-question", level=2)

doc.add_heading("SQ1: What types of platform strategies are observed in biotech firms?", level=3)
doc.add_paragraph(
    "Four platform types emerge from the corpus. Technology platforms are the most clearly evidenced (6 papers), including biomanufacturing, drug discovery, and agbiotech breeding platforms. Multi-sided platforms appear in only one paper, covering open data platforms mediating between data suppliers and consumers. R&D/open innovation platforms are functionally described across several papers involving consortia, data-sharing infrastructure, and cross-firm collaboration networks. Digital/data platforms (blockchain, bio-search, genomics repositories) emerge as a distinct contemporary category. Notably, modular product platforms, identified in the protocol as a distinct strategy type, find almost no treatment in the corpus."
)

doc.add_heading("SQ2: What firm-level outcomes are associated with platform strategies?", level=3)
doc.add_paragraph(
    "Innovation output is the most frequently measured outcome (35+ papers) but is rarely operationalized rigorously. Partnership and alliance formation is the second most common, with evidence that platform strategies facilitate partnerships, particularly during crisis conditions. However, at least one high-quality paper (MMAT 4.5) argues that existing collaboration is structurally inefficient. The strongest quantitative evidence shows negative long-run financial performance for biotech IPO firms (event study, MMAT 4.5), and no paper demonstrates that platform-adopting firms outperform non-platform peers on financial metrics. Competitive advantage is frequently cited but rarely measured with observable indicators. Survival is measured in four papers with no robust positive effects. Time-to-market does not appear as a measured outcome in any included paper."
)

doc.add_heading("SQ3: Under what conditions do platform strategies create or destroy value?", level=3)
doc.add_paragraph(
    "The dominant finding is that platform effects are conditional, not universal. Six key moderators emerge: (1) Ecosystem maturity: the highest-quality paper (MMAT 5.0) establishes that platform value depends on surrounding digital and institutional ecosystem maturity. (2) Governance design: open data platforms require explicit governance structures. (3) Regulatory and payor environments: regulatory hurdles can neutralize platform efficiency gains. (4) Firm dynamic capabilities: value creation depends on the firm's ability to sense, seize, and reconfigure resources. (5) Geography: emerging markets develop distinct platform models. (6) Crisis conditions: platform infrastructure is most visibly valuable during health crises but fragility is simultaneously exposed."
)

doc.add_heading("SQ4: What boundary conditions and trade-offs are documented?", level=3)
doc.add_paragraph(
    "Six trade-offs are documented: (1) Openness vs. appropriability: tension between open innovation and IP protection. (2) Scalability vs. customization: overly rigid platforms may fail in therapeutic contexts requiring personalization. (3) Global reach vs. local access: technology platforms create access gaps in LMICs. (4) Short-term returns vs. long-term value: biotech platform investment consistently underperforms in the short-to-medium run. (5) Collaboration breadth vs. coordination costs: maximizing partner diversity increases governance complexity. (6) Technological ambition vs. regulatory feasibility: advanced platforms face regulatory environments not yet adapted to platform-based development."
)

doc.add_page_break()

# ===================== 3.4 Included Studies Table =====================
doc.add_heading("3.4 Included Studies", level=2)

# Compact table: ID, Method, Platform Type, Effect Direction, Key Finding
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Paper ID", "Method", "Platform Type", "Effect", "Key Finding"]):
    table.rows[0].cells[i].text = h
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(8)

for paper in papers:
    row = table.add_row()
    pid = paper["id"]
    if len(pid) > 30:
        pid = pid[-30:]
    row.cells[0].text = pid
    row.cells[1].text = (paper.get("method") or "n/a")[:20]
    row.cells[2].text = (paper.get("platform_type") or "n/a")[:15]
    row.cells[3].text = (paper.get("effect_direction") or "n/a")[:12]
    summary = paper.get("effect_summary") or ""
    row.cells[4].text = summary[:80] + ("..." if len(summary) > 80 else "")
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(7)

doc.add_page_break()

# ===================== 3.5 Quality Assessment Table =====================
doc.add_heading("3.5 Quality Assessment (MMAT 2018)", level=2)

table = doc.add_table(rows=1, cols=5)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Paper ID", "MMAT Category", "Total Score", "Score Band", "Notes"]):
    table.rows[0].cells[i].text = h
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(8)

for pid, q in sorted(quality.items(), key=lambda x: -x[1]["total"]):
    row = table.add_row()
    display_id = pid if len(pid) <= 30 else pid[-30:]
    row.cells[0].text = display_id
    row.cells[1].text = q["method_category"]
    row.cells[2].text = str(q["total"])
    t = q["total"]
    band = "Excellent" if t >= 4 else "Good" if t >= 3 else "Adequate" if t >= 2 else "Weak" if t >= 1 else "Very Weak"
    row.cells[3].text = band
    row.cells[4].text = (q.get("notes") or "")[:60]
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(7)

doc.add_paragraph()
doc.add_paragraph(
    "Quality score distribution: Excellent (4.0-5.0): 10 papers; Good (3.0-3.9): 10 papers; Adequate (2.0-2.9): 10 papers; Weak (1.0-1.9): 12 papers; Very Weak (0.0): 10 papers. Mean score: 2.17/5.0."
)

doc.add_page_break()

# ===================== 4. DISCUSSION =====================
doc.add_heading("4. Discussion", level=1)

doc.add_paragraph(
    "This systematic review synthesized 52 papers on the effects of platform strategies in biotech companies. The overarching finding is that platform effects are conditional rather than universally positive or negative. This contrasts with the more optimistic framing common in the broader platform strategy literature, where network effects and winner-take-all dynamics are often presented as inherently value-creating."
)
doc.add_paragraph(
    "The evidence is strongest for technology platforms in biomanufacturing and drug discovery, where incremental production improvements are well documented. However, the translation from technical innovation to firm-level financial performance remains undemonstrated. The single rigorous quantitative study (event study of 319 biotech IPOs, MMAT 4.5) shows negative long-run financial performance across the sector, providing a sobering counterpoint to the positive innovation narratives."
)
doc.add_paragraph(
    "Three contradictions in the literature deserve attention. First, partnership formation is simultaneously described as a benefit and a structural failure of current platforms. Second, positive innovation effects coexist with negative financial effects, suggesting a disconnect between technical and commercial value. Third, ecosystem openness is framed as enabling by lower-quality papers but as governance-dependent by the highest-quality studies."
)
doc.add_paragraph(
    "The dominance of qualitative case studies (48 of 52 papers) severely limits causal inference. The field needs large-sample quantitative studies that directly compare platform-adopting and non-platform biotech firms on measurable outcomes."
)

# ===================== 5. RESEARCH GAPS =====================
doc.add_heading("5. Research Gaps", level=1)

gaps = [
    "No study directly compares platform-adopting vs. non-platform biotech firms on financial outcomes. A panel dataset tracking revenue, R&D productivity, and survival over 10+ years would substantially advance the field.",
    "Modular product platforms are entirely absent. Future research should investigate modular architectures in therapeutics (e.g., bispecific antibody platforms, modular CAR-T constructs).",
    "Time-to-market is not measured in any included paper despite being a major value driver in biotech.",
    "Multi-sided platform dynamics are under-studied, with only one paper examining this configuration in life sciences.",
    "LMIC biotech platform dynamics need dedicated evidence. Two papers are insufficient for generalizable claims.",
    "Platform governance design has no empirical benchmarks. Comparative case studies of governance structures in platform consortia are needed.",
    "The interaction between firm size and platform strategy is not examined in any included paper.",
    "Therapeutic area as a moderator of platform applicability is unaddressed.",
    "Survival as an outcome is poorly measured. A discrete-time hazard model using platform strategy as a time-varying covariate would address this.",
    "The evolution of platform strategies across the firm lifecycle is not tracked. Longitudinal case studies from founding through exit would provide foundational evidence."
]

for i, gap in enumerate(gaps, 1):
    doc.add_paragraph(f"{i}. {gap}")

# ===================== 6. LIMITATIONS =====================
doc.add_heading("6. Limitations", level=1)

doc.add_paragraph(
    "This review has several limitations that should be considered when interpreting the findings:"
)
limitations = [
    "Language restriction: Only English-language papers were included, potentially excluding relevant research published in other languages, particularly from Asian biotech markets.",
    "Database scope: Semantic Scholar returned zero results due to API limitations. PubMed and Scopus were not searched. The 2,000-record cap on OpenAlex and CrossRef means some relevant papers may have been missed.",
    "Screening reliability: A single automated screener was used rather than dual independent screening with inter-rater reliability assessment, as recommended by PRISMA.",
    "Data extraction limitations: Heuristic-based extraction from full texts introduces noise. Author metadata was missing for the majority of papers, preventing author-year citation in the synthesis.",
    "Quality assessment: MMAT scoring was applied via keyword heuristics rather than expert judgment, potentially over- or under-scoring papers.",
    "Publication bias: No formal assessment of publication bias (e.g., funnel plot) was conducted, as the predominantly qualitative corpus does not lend itself to standard methods.",
    "Full-text retrieval: Only 53 of 92 included papers had accessible full text (58% retrieval rate). The 39 paywalled papers represent a potentially non-random subsample loss."
]
for lim in limitations:
    doc.add_paragraph(lim, style='List Bullet')

# ===================== 7. CONCLUSIONS =====================
doc.add_heading("7. Conclusions", level=1)

doc.add_paragraph(
    "Platform strategies in biotechnology produce conditional rather than universal benefits. The literature is dominated by qualitative case studies with limited causal evidence. Technology platforms and R&D/open innovation platforms are the most observed types, while multi-sided and modular platforms are largely absent from the biotech literature. The strongest quantitative evidence points to negative long-run financial performance for biotech firms, though this is sector-wide rather than platform-specific. Key moderators of platform value include ecosystem maturity, governance design, regulatory environment, and firm dynamic capabilities. Six distinct trade-offs are documented, most notably between platform openness and IP appropriability. Ten specific research gaps are identified that, if addressed, would substantially strengthen the evidence base for strategic decision-making in the biotech sector."
)

# ===================== 8. REFERENCES =====================
doc.add_heading("8. References", level=1)
doc.add_paragraph(
    "Note: Due to metadata limitations in the retrieved corpus, full bibliographic details are not available for all included papers. Papers are referenced by their DOI or database identifier throughout this report."
)
doc.add_paragraph()

# List papers with DOIs as references
for paper in sorted(papers, key=lambda p: p.get("doi") or p["id"]):
    doi = paper.get("doi")
    pid = paper["id"]
    year = paper.get("year") or "n.d."
    method = paper.get("method") or "unspecified method"
    if doi:
        doc.add_paragraph(f"{pid}. ({year}). [{method}]. https://doi.org/{doi}", style='List Bullet')
    else:
        doc.add_paragraph(f"{pid}. ({year}). [{method}]. [No DOI available]", style='List Bullet')

# Save
out_path = BASE / "output" / "slr_report.docx"
doc.save(str(out_path))
print(f"Report saved to {out_path}")
