#!/usr/bin/env python3
"""Build the PRISMA 2020 SLR report as a .docx from synthesis + extraction + quality data."""
import json
from pathlib import Path
from collections import Counter

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = Path("/home/npdrpi1/ClaudeAgents/slr-agent-master")
EXTRACT = ROOT / "data/05_extracted/extraction.jsonl"
QUALITY = ROOT / "data/05_extracted/quality.jsonl"
COUNTS = ROOT / "output/prisma_counts.json"
OUT = ROOT / "output/slr_report.docx"


def load_jsonl(p):
    rows = []
    for line in p.read_text().splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            rows.append(json.loads(line))
        except json.JSONDecodeError:
            continue
    return rows


extraction = load_jsonl(EXTRACT)
quality = load_jsonl(QUALITY)
qmap = {q["id"]: q for q in quality}
counts = {}
for line in COUNTS.read_text().splitlines():
    if line.strip():
        o = json.loads(line)
        counts[o["phase"]] = o["n"]

# ---- Derived statistics ----
extracted_ok = [e for e in extraction if e.get("extraction_status") == "extracted"]

# included final = MMAT >= 2
included = sorted(
    [e for e in extraction if qmap.get(e["id"], {}).get("total", 0) >= 2],
    key=lambda e: -qmap.get(e["id"], {}).get("total", 0),
)

method_counts = Counter(
    (e.get("method") or "not stated") for e in extracted_ok
)
lens_counts = Counter(
    (e.get("theoretical_lens") or "none stated") for e in extracted_ok
)
geo_counts = Counter()
for e in extracted_ok:
    g = (e.get("sample") or {}).get("geography")
    if g:
        geo_counts[g] += 1
ptype_counts = Counter(
    (e.get("platform_type") or "not stated") for e in extracted_ok
)
dir_counts = Counter(
    (e.get("effect_direction") or "not stated") for e in extracted_ok
)
qtotals = Counter(qmap.get(e["id"], {}).get("total", 0) for e in extracted_ok)

identified_total = sum(v for k, v in counts.items() if k.startswith("identified_"))

# Curated on-topic finding-bearing records (from synthesis) for the included table.
# Maps id -> (platform_label, key_finding) drawn from synthesis themes (<=15 words each).
FINDING = {
    "10.1186_s13731-015-0027-3": ("Technology / R&D", "New approaches needed to enhance pipeline productivity, valuation, risk management (conditional)."),
    "10.1371_journal.pone.0243813": ("Therapeutics IPO firms", "Biotech IPO firms underperformed paired non-biotech controls financially (negative)."),
    "10.1186_1478-4505-9-37": ("Technology platform", "Coded negative on innovation/performance/alliance outcomes."),
    "https___openalex.org_W3153452158": ("Ecosystem / business model", "Critiques inefficient collaboration and an overly-financialized business model."),
    "10.3389_fbloc.2020.586525": ("Technology (platform theory)", "Platform-theory framing of innovation/performance/partnership."),
    "10.1186_1471-2105-15-s1-s2": ("Data-integration platform", "Life-science data capture/integration tools; mixed partnership/growth effects."),
    "10.1371_journal.pone.0276204": ("Multi-sided data platform", "Open data platforms mediate user data supply and demand (conditional)."),
    "10.1186_s44398-025-00004-7": ("Discovery (bioprospecting)", "Bioprospecting discovery platform; innovation/partnership effects conditional."),
    "10.1186_1472-698x-10-s1-s1": ("Emerging-market model", "Emerging markets developed lower-cost models and technological innovations."),
    "10.3389_fbloc.2025.1510429": ("Ecosystem platform", "Innovation/partnership/survival outcomes coded conditional."),
    "10.35808_ersj_2088": ("Collaboration network", "Documents collaborative program creation and development."),
    "EPMC_41095161_MED": ("Ecosystem (agbiotech)", "Off-topic extraction; nominally conditional ecosystem effects."),
    "https___openalex.org_W2904104880": ("Knowledge / IP substrate", "Knowledge treated as a managed critical organizational asset (mixed)."),
    "https___openalex.org_W3171428381": ("IP / patent strategy", "Examines patent pledges as a strategy (conditional)."),
    "https___openalex.org_W188916526": ("Open-innovation platform", "Post-genomic open innovation sustains growth in knowledge-driven competition."),
    "10.1007_10_2018_59": ("Bioprocessing technology", "Process optimization raised antibody titres to 3-8 g/L (positive)."),
    "10.3846_jbem.2019.6880": ("Business-model / circular economy", "Links platform/business-model design to the circular economy."),
    "10.1007_s00122-025-05060-1": ("Breeding platform (agbiotech)", "Framed as transformation engine for data-driven breeding innovation (positive)."),
    "EPMC_42190778_MED": ("Transaction-cost framing", "Innovation/performance effects coded conditional."),
    "10.5912_jcb1280": ("Competitive-advantage case", "Competitive-advantage outcome coded conditional."),
    "10.3389_fbioe.2025.1688121": ("Biosensor technology", "Off-topic extraction (nanophotonic biosensors); no usable effect."),
    "10.1017_s109285292200092x": ("Open-innovation (brain health)", "Open innovation framed as key to advancing the field (conditional)."),
    "10.1007_s44337-025-00313-w": ("Global-health technology", "COVID-19 exposed fragility of developing-world health systems."),
    "https___openalex.org_W780003724": ("Innovation/performance (other)", "Innovation and performance outcomes; direction not stated."),
}

# ============================ DOCUMENT ============================
doc = Document()

# base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def para(text, italic=False, bold=False, size=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t

# ---------------- TITLE PAGE ----------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("The Effect of Platform Strategies in Biotechnology Companies")
tr.bold = True
tr.font.size = Pt(20)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("A Systematic Literature Review (PRISMA 2020)")
sr.font.size = Pt(14)
sr.italic = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Review date: 17 June 2026\n").font.size = Pt(11)
meta.add_run(f"Databases searched: OpenAlex, CrossRef, EuropePMC, Semantic Scholar\n").font.size = Pt(11)
meta.add_run(f"Records identified: {identified_total:,}  |  Studies included in synthesis: {len(included)}").font.size = Pt(11)

doc.add_paragraph()
rq = doc.add_paragraph()
rq.alignment = WD_ALIGN_PARAGRAPH.CENTER
rqr = rq.add_run('Research question: "What effect do platform strategies have in biotech companies?"')
rqr.italic = True
rqr.font.size = Pt(12)

doc.add_page_break()

# ---------------- ABSTRACT ----------------
heading("Abstract", level=1)

para("Background. ", bold=True, space_after=0)
para(
    "Biotechnology firms increasingly frame their strategy around platforms — reusable "
    "technology and R&D capabilities, multi-sided data infrastructures, and ecosystem "
    "business models — rather than single assets. Whether these platform strategies actually "
    "improve firm-level outcomes is contested, and no prior review has consolidated the evidence "
    "specific to commercial biotech firms.",
)
para("Objectives. ", bold=True, space_after=0)
para(
    "To systematically identify, appraise, and synthesise empirical and conceptual evidence on "
    "the types, outcomes, and boundary conditions of platform strategies in biotechnology companies.",
)
para("Methods. ", bold=True, space_after=0)
para(
    f"Following a pre-registered PRISMA 2020 protocol, four databases (OpenAlex, CrossRef, EuropePMC, "
    f"Semantic Scholar) were searched for English-language records published 2005 or later. "
    f"{identified_total:,} records were identified; after deduplication ({counts['after_dedup']:,}) and "
    f"keyword pre-filtering ({counts['after_prefilter']:,}), titles and abstracts were screened. "
    f"Full texts were retrieved ({counts['full_text_retrieved']}), data were extracted, and studies were "
    f"appraised with the Mixed Methods Appraisal Tool (MMAT 2018). Findings were combined by thematic synthesis.",
)
para("Results. ", bold=True, space_after=0)
para(
    f"{len(included)} studies met the quality threshold (MMAT ≥ 2) and entered the synthesis. The corpus "
    "is methodologically thin and dominated by ecosystem-lens case studies; effect direction is unreported "
    "for roughly two-thirds of records. Where stated, the modal coded effect is conditional, not positive. "
    "Three de-facto platform conceptions recur: collaboration/open-innovation platforms, technology/R&D "
    "discovery platforms, and business-model/ecosystem platforms. Crucially, the two strongest quantitative "
    "studies coded negative on financial and survival outcomes, while positive claims clustered in weaker "
    "qualitative work — a quality-direction confound. No study tested a named moderator (firm size, "
    "therapeutic area, financing stage, ecosystem maturity).",
)
para("Conclusions. ", bold=True, space_after=0)
para(
    "The current evidence is too sparse, uneven in quality, and inconsistent to support a confident claim "
    "that platform strategies improve biotechnology firm performance. Platforms may raise innovation and "
    "collaboration activity without translating into superior financial returns. Rigorous, design-justified, "
    "moderator-testing studies are the priority gap.",
)

doc.add_page_break()

# ---------------- BACKGROUND ----------------
heading("1. Background", level=1)
para(
    "Platform strategy has become one of the dominant explanations for competitive advantage in "
    "technology-intensive industries. In biotechnology specifically, firms are increasingly organised "
    "around reusable scientific and technological capabilities — discovery engines, data infrastructures, "
    "and modular R&D capacity — that can be redeployed across multiple programmes rather than around a "
    "single therapeutic or product asset. Proponents argue that such platforms lower the marginal cost of "
    "each new programme, attract partners and capital, and create defensible ecosystems around the firm."
)
para(
    "This matters because biotech is capital-intensive, long-cycle, and historically loss-making at the "
    "sector level; one record in this corpus notes accumulated sector losses since the mid-1970s. If "
    "platform strategies genuinely improve pipeline productivity, valuation, partnership formation, or "
    "survival, they would offer a route out of that structural unprofitability. If they do not, the "
    "platform narrative may be over-sold."
)
para(
    "Prior platform-strategy reviews tend to be cross-industry or focused on digital consumer platforms, "
    "leaving the biotech-specific evidence scattered and unconsolidated. This review addresses that gap by "
    "asking: what effect do platform strategies have in biotech companies? Four sub-questions follow the "
    "protocol: (SQ1) what types of platform strategies are observed; (SQ2) what firm-level outcomes are "
    "associated with them; (SQ3) under what conditions do they create or destroy value; and (SQ4) what "
    "boundary conditions and trade-offs are documented."
)

# ---------------- METHODS ----------------
heading("2. Methods", level=1)
para(
    "This review was conducted and reported in accordance with the PRISMA 2020 statement and a "
    "pre-specified protocol (v0.1). The methods below restate that protocol so that the search is replicable."
)

heading("2.1 PICOC framing", level=2)
add_table(
    ["Element", "Specification"],
    [
        ["Population", "Biotechnology firms (therapeutics, diagnostics, agbiotech, industrial biotech); pure pharma incumbents only as comparison groups."],
        ["Intervention", "Adoption/use of a platform strategy (technology, multi-sided, modular product, or R&D platform)."],
        ["Comparison", "Non-platform / pipeline / single-asset strategies, or pre/post adoption."],
        ["Outcome", "Innovation output, financial performance, valuation, partnership/alliance formation, survival, time-to-market, ecosystem dynamics."],
        ["Context", "Any geography; commercial biotech firms, public or private."],
    ],
    widths=[1.3, 5.0],
)

heading("2.2 Databases and search dates", level=2)
para(
    "Four databases were searched: OpenAlex (primary, broadest recall), Semantic Scholar (cross-validation "
    "and abstract summaries), EuropePMC (biomedical depth and full-text retrieval), and CrossRef (metadata "
    "reconciliation and DOI canonicalisation). The date filter was 2005-01-01 to the search date (2026)."
)

heading("2.3 Search query", level=2)
para(
    "The base free-text Boolean query (translated per database by the query-builder) was three AND-linked "
    "concept blocks:"
)
q = doc.add_paragraph()
qr = q.add_run(
    '( "platform strateg*" OR "platform business model" OR "platform ecosystem" OR "multi platform" OR '
    '"platform architecture" OR "innovation platform" OR "technology platform" OR "digital platform" OR '
    '"platform-based" OR "platform competition" OR "platform governance" OR "knowledge platform" OR '
    '"modular platform" )\nAND\n( "biotechnology" OR "biotech" OR "biopharmaceutical" OR "biopharma" OR '
    '"life sciences" OR "biotech firm*" OR "biotech startup*" OR "biotech compan*" )\nAND\n'
    '( "innovation performance" OR "competitive advantage" OR "value creation" OR "open innovation" OR '
    '"firm performance" OR "strategic alliance*" OR "drug development" OR "collaboration" OR "scalabilit*" )'
)
qr.font.name = "Consolas"
qr.font.size = Pt(8.5)

heading("2.4 Inclusion and exclusion criteria", level=2)
para("Included if all hold:", bold=True, space_after=0)
for c in [
    "Peer-reviewed journal article, conference paper, or established working paper",
    "Empirical (qualitative, quantitative, mixed) OR substantive conceptual/theory paper",
    "Explicitly addresses platform strategy AND a biotech / life-science context",
    "Published 2005 or later; English language; full text obtainable",
]:
    doc.add_paragraph(c, style="List Bullet")
para("Excluded if any hold:", bold=True, space_after=0)
for c in [
    "Editorial, book review, news item, or blog post",
    "Platform discussed only as background, or biotech used only as an illustrative example",
    '"Platform" used in a purely technical/laboratory sense (e.g. "PCR platform") without strategic content',
    "Pharma-only with no biotech firms in the sample; pre-2005; non-English; full text inaccessible after two attempts",
]:
    doc.add_paragraph(c, style="List Bullet")
para(
    "Borderline records were labelled MAYBE rather than excluded (ambiguous strategic content, partial biotech "
    "focus, or unclear methodology) and passed to a second-pass screen for adjudication."
)

heading("2.5 Screening, extraction, and quality assessment", level=2)
para(
    "After deduplication and a keyword pre-filter, records were screened on title and abstract in batches. "
    "Data were extracted per paper into a structured schema (bibliographic fields, theoretical lens, method, "
    "sample, platform definition and type, outcomes measured, effect direction and magnitude, boundary "
    "conditions, and stated gaps). Quality was appraised with the MMAT 2018 (five items, scored 0–5). Per the "
    "protocol, MMAT scores are reported transparently; for this synthesis a threshold of MMAT ≥ 2 was applied "
    "to define the final included set, so that the lowest-information records do not drive conclusions."
)

heading("2.6 Synthesis approach", level=2)
para(
    "Findings were combined by thematic synthesis (Thomas & Harden, 2008): line-by-line coding of on-topic "
    "extracted text into descriptive themes, then analytical themes mapped back to the four sub-questions. "
    "A documented data-quality caveat applies: a substantial fraction of extracted effect-summary fields "
    "contained extraction artefacts (table-of-contents fragments, citation metadata, off-topic first "
    "sentences). Such records were treated as low-information and excluded from claims about effect direction. "
    "Only ~12–15 records carried substantively on-topic content; this conservatism is itself a headline finding."
)

# ---------------- RESULTS ----------------
heading("3. Results", level=1)

heading("3.1 Study selection (PRISMA flow)", level=2)
para(
    f"Database searches identified {identified_total:,} records (OpenAlex {counts.get('identified_openalex',0):,}, "
    f"CrossRef {counts.get('identified_crossref',0):,}, EuropePMC {counts.get('identified_epmc',0):,}, "
    f"Semantic Scholar {counts.get('identified_s2',0):,}). After duplicate removal {counts['after_dedup']:,} records "
    f"remained, and {counts['after_prefilter']:,} passed the keyword pre-filter and were screened on title and "
    f"abstract. Screening yielded {counts['screened_include']} INCLUDE, {counts['screened_maybe']:,} MAYBE, and "
    f"{counts['screened_exclude']:,} EXCLUDE decisions. Full text was retrieved for {counts['full_text_retrieved']} "
    f"records; {counts['extracted']} were successfully extracted (one returned insufficient text). Applying the "
    f"MMAT ≥ 2 quality threshold left {len(included)} studies in the final synthesis. The flow is shown in "
    f"output/prisma_flow.svg (Figure 1)."
)
fig = doc.add_paragraph()
fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fig.add_run("Figure 1. PRISMA 2020 flow diagram (see output/prisma_flow.svg).")
fr.italic = True
fr.font.size = Pt(9)

heading("3.2 Study characteristics", level=2)
para(
    "The corpus is descriptively thin. Year metadata is unreliable (several records carry implausible future "
    "years and many period fields contain ISSNs rather than dates); where populated, years cluster 2010–2026. "
    "Geography is mostly unrecorded and, where stated, skews to the UK, USA, and Europe. Methods are dominated "
    "by case study / qualitative work, with a minority of quantitative-descriptive regressions and a large set "
    "of records with no stated method. Ecosystem theory is the most common explicit lens, but the single most "
    "common state is no stated lens at all. Median MMAT total across extracted records is approximately 1."
)

para("Table 1. Method distribution (extracted records).", italic=True, size=9, space_after=2)
add_table(["Method", "n"], [[k, v] for k, v in method_counts.most_common()], widths=[3.0, 1.0])

para("Table 2. Theoretical lens distribution (extracted records).", italic=True, size=9, space_after=2)
add_table(["Theoretical lens", "n"], [[k, v] for k, v in lens_counts.most_common()], widths=[3.0, 1.0])

para("Table 3. Reported geography (extracted records with a stated geography).", italic=True, size=9, space_after=2)
add_table(["Geography", "n"], [[k, v] for k, v in geo_counts.most_common()], widths=[3.0, 1.0])

para("Table 4. Platform type as coded (extracted records).", italic=True, size=9, space_after=2)
add_table(["Platform type", "n"], [[k, v] for k, v in ptype_counts.most_common()], widths=[3.0, 1.0])

para("Table 5. MMAT total score distribution (extracted records).", italic=True, size=9, space_after=2)
add_table(["MMAT total (0–5)", "n"], [[str(k), v] for k, v in sorted(qtotals.items())], widths=[3.0, 1.0])

para("Table 6. Coded effect direction (extracted records).", italic=True, size=9, space_after=2)
add_table(["Effect direction", "n"], [[k, v] for k, v in dir_counts.most_common()], widths=[3.0, 1.0])

heading("3.3 Findings by theme", level=2)
para(
    "Thematic synthesis produced seven descriptive themes and four analytical themes mapped to the "
    "sub-questions. The themes below rest only on records whose extracted text plausibly bears on platform "
    "strategy in a biotech/life-science firm; thin or noisy evidence is flagged."
)

theme_blocks = [
    ("Theme 1 — Platforms as collaboration / partnership engines",
     "Several records frame platforms primarily as mechanisms for forming alliances, pooling data, and "
     "orchestrating multi-party collaboration rather than as product architectures. Open-innovation and "
     "partnership framing recurs across post-genomic open-innovation and pandemic-partnership records. The "
     "evidence is broad but mostly low-quality."),
    ("Theme 2 — Platforms as business-model / value-creation reconfiguration",
     "A cluster treats the platform as a redesign of value creation and capture, including framing a business "
     "model as the design of value-creation and capture mechanisms (coded conditional) and tying "
     "platform/business-model work to the circular economy. Anchored by lower-quality records (MMAT 1–2)."),
    ("Theme 3 — Technology / R&D platforms and pipeline productivity",
     "A small but better-anchored set links platform strategy to R&D pipeline productivity, valuation, and "
     "risk management. The highest-quality record in the corpus (MMAT 5, resource-based view, regression) "
     "argues new creative approaches are needed to enhance pipeline productivity, valuation, and risk "
     "management, and codes the effect conditional. This is the best-anchored theme."),
    ("Theme 4 — Knowledge management and IP/patent strategy as platform substrate",
     "Records here treat knowledge assets and patent strategy as the substrate of platform value: knowledge "
     "as a managed critical asset (coded mixed) and patent pledges as a deliberate strategy (coded conditional). "
     "Evidence is tentative (1–2 anchors)."),
    ("Theme 5 — Platforms and financial / survival underperformance",
     "A high-quality counter-current questions whether platform-era biotech delivers financially. The strongest "
     "case (MMAT 4, regression) compared 319 therapeutics-focused biotech IPOs (1997–2016) against paired "
     "non-biotech controls and is coded negative on performance, valuation, and survival; a second MMAT 4 "
     "regression also codes negative. A qualitative record notes large accumulated sector losses. This theme is "
     "thin in number but evidentially heavy because the designs are the strongest in the corpus."),
    ("Theme 6 — Ecosystem maturity, mission-orientation, and systemic problems",
     "Ecosystem-lens records situate platform value in the surrounding system. An MMAT 4 case study critiques "
     "inefficient collaboration and an overly-financialized business model, while agbiotech breeding-platform "
     "records frame platforms as transformation engines for data-driven, demand-responsive innovation."),
    ("Theme 7 — Emerging-market and global-health platform models (low information)",
     "A thin cluster on emerging-market / global-health business models exists but is largely extraction noise; "
     "one record notes emerging markets developed lower-cost models and technological innovations. Weakly "
     "evidenced and partly off-topic."),
]
for h, body in theme_blocks:
    heading(h, level=3)
    para(body)

heading("3.4 Mapping to the sub-questions", level=3)
sq = [
    ("SQ1 — Types of platform strategy.",
     "The corpus does not cleanly use the protocol's four-part typology. Most records collapse to \"other\" or "
     "are untyped; only a handful are coded technology and exactly one multi-sided (open data platforms). No "
     "record is coded as a modular product platform. Reading through the noise, three de-facto conceptions "
     "appear: collaboration/open-innovation platforms; technology/R&D discovery platforms; and "
     "business-model/ecosystem platforms."),
    ("SQ2 — Firm-level outcomes.",
     "Effect direction is absent for roughly two-thirds of records. Among those with a stated direction the "
     "modal value is conditional, not positive. Positive claims (5) are mostly qualitative and lower-quality "
     "and centre on innovation/partnership outcomes; the two negative codes are both quantitative and "
     "high-quality and centre on financial/survival outcomes. The strongest designs lean "
     "conditional-to-negative on hard financial outcomes — a quality-direction confound."),
    ("SQ3 — Conditions of value creation/destruction.",
     "This is the weakest-supported sub-question. Boundary-condition fields were essentially empty, so explicit "
     "moderators (firm size, therapeutic area, financing stage, ecosystem maturity) are undocumented. Only "
     "indirect signals exist: ecosystem maturity/coordination as an enabler, pipeline productivity as the "
     "lever for R&D platforms, and IPO/financing stage as a context where platform-era biotech underperformed."),
    ("SQ4 — Boundary conditions and trade-offs.",
     "Almost nothing is documented explicitly. Inferable, untested trade-offs include openness vs "
     "appropriability (open-innovation and patent-pledge framings), collaboration value vs financialization, "
     "and scientific productivity vs financial return (rising technical yields alongside sector-level losses). "
     "These are analyst inferences across papers, offered as hypotheses for future testing."),
]
for h, body in sq:
    heading(h, level=4)
    para(body)

heading("3.5 Included studies", level=2)
para(
    "Table 7 lists the studies that met the MMAT ≥ 2 threshold and entered the synthesis, ordered by quality "
    "score. Author and year metadata were largely missing or unreliable in the source records, so studies are "
    "identified by their stable record ID/DOI. Findings are summarised from the synthesis, not from raw "
    "extraction fields.",
)
para("Table 7. Included studies (MMAT ≥ 2).", italic=True, size=9, space_after=2)
inc_rows = []
for e in included:
    q = qmap.get(e["id"], {})
    label, finding = FINDING.get(e["id"], ((e.get("platform_type") or "not stated"), "On-topic content limited; see synthesis."))
    inc_rows.append([
        e["id"],
        (e.get("method") or "n/s"),
        label,
        (e.get("effect_direction") or "n/s"),
        q.get("total", 0),
        finding,
    ])
add_table(
    ["Record ID / DOI", "Method", "Platform type", "Effect", "MMAT", "Key finding"],
    inc_rows,
    widths=[1.7, 0.8, 1.2, 0.7, 0.5, 2.4],
)

heading("3.6 Quality assessment", level=2)
para(
    "Table 8 reports the MMAT appraisal for the included studies. Most low scores in the wider corpus were "
    "driven by failing item q2 (appropriate/justified design) and item q4 (coherence / risk-of-bias control). "
    "Only the single MMAT 5 record satisfied all five items.",
)
para("Table 8. MMAT quality assessment of included studies.", italic=True, size=9, space_after=2)
q_rows = []
for e in included:
    q = qmap.get(e["id"], {})
    s = q.get("scores", {})
    q_rows.append([
        e["id"],
        q.get("method_category", "n/s"),
        s.get("q1", ""), s.get("q2", ""), s.get("q3", ""), s.get("q4", ""), s.get("q5", ""),
        q.get("total", 0),
    ])
add_table(
    ["Record ID / DOI", "MMAT category", "q1", "q2", "q3", "q4", "q5", "Total"],
    q_rows,
    widths=[2.0, 1.3, 0.4, 0.4, 0.4, 0.4, 0.4, 0.5],
)

# ---------------- DISCUSSION ----------------
heading("4. Discussion", level=1)
para(
    "Taken together, the evidence does not support a confident claim that platform strategies improve "
    "biotechnology firm performance. The literature is heterogeneous in what it even means by a \"platform\" "
    "— collaboration infrastructure, reusable R&D/discovery capability, or business-model/ecosystem redesign "
    "— and it rarely converges on a sign for the effect. Where an effect direction is stated at all, the modal "
    "verdict is conditional."
)
para(
    "Where the evidence is comparatively solid, it points to nuance rather than enthusiasm. The best-anchored "
    "theme (technology/R&D platforms, including the only MMAT 5 record) frames platform value as contingent on "
    "pipeline productivity and risk management. The most consequential and best-designed finding is negative: "
    "the two MMAT 4 quantitative studies code platform-era biotech as underperforming, including a matched "
    "comparison of 319 therapeutics IPOs against non-biotech controls."
)
para(
    "This produces a quality-direction confound that should temper any optimistic reading. Positive claims "
    "cluster in weaker, qualitative, innovation- and partnership-framed work, while the negative signals come "
    "from the strongest designs. A plausible reconciliation is that platforms may raise innovation activity "
    "and collaboration without translating into superior risk-adjusted financial returns — a pattern "
    "consistent with the sector's documented accumulated losses. Higher MMAT scores also sit with the more "
    "critical, less celebratory readings of the platform model."
)
para(
    "The evidence is thinnest exactly where it matters most for practice. No study tested a named moderator, "
    "so SQ3 — when platforms create versus destroy value — is effectively unanswered. Boundary conditions and "
    "trade-offs are asserted but never measured. Causal and financial designs (event studies, panel-causal, "
    "matched-difference) are almost entirely absent beyond the two negative-leaning regressions."
)

heading("4.1 Research gaps", level=2)
for g in [
    "No tested moderators: not one record supplies an interaction effect for firm size, therapeutic area, financing stage, or ecosystem maturity.",
    "Outcome direction missing at scale: roughly two-thirds of records lack any coded effect direction.",
    "Causal/financial designs absent beyond two negative-leaning regressions, despite SQ2 listing valuation, time-to-market, and survival.",
    "Typology under-specified: technology vs multi-sided vs modular vs R&D platforms are rarely distinguished, and modular product platforms are entirely absent.",
    "Innovation-to-financial-value translation: the implied disconnect between rising technical/innovation output and weak financial returns is untested directly.",
    "Geographic and stage imbalance: evidence is thin everywhere and skewed to UK/US/Europe.",
    "Boundary conditions and trade-offs (openness vs appropriability; value creation vs financialization) are unmeasured.",
    "Reporting and appraisal quality: median MMAT ~1, with systematic failure of items q2 and q4; the field needs pre-registered, design-justified studies.",
]:
    doc.add_paragraph(g, style="List Bullet")

heading("4.2 Limitations", level=2)
para(
    "Several limitations qualify these conclusions. (1) Language: only English-language records were eligible, "
    "excluding potentially relevant non-English work. (2) Database scope: four databases were searched and "
    "Semantic Scholar returned no records in this run, so recall depends largely on OpenAlex, CrossRef, and "
    "EuropePMC; grey literature and patents were not systematically covered. (3) Screening reliability: the "
    "large MAYBE pool and batch screening mean inter-rater reliability was not formally established as in a "
    "fully dual-screened review, so selection error cannot be ruled out. (4) Data quality: a substantial share "
    "of extracted fields were artefacts, which constrained the synthesis to ~12–15 substantively on-topic "
    "records and made year, geography, and platform-type statistics unreliable. (5) Publication and reporting "
    "bias: with effect direction unreported in most records and appraisal quality low, both publication bias "
    "and selective reporting are plausible and could not be tested (e.g. via a funnel plot)."
)

# ---------------- CONCLUSIONS ----------------
heading("5. Conclusions", level=1)
para(
    "Across the studies we could read, \"platform strategy\" in biotech means several different things — ways "
    "to collaborate and share data, reusable R&D and discovery technologies, and new business-model designs "
    "— and the studies rarely agree on whether these strategies pay off. Lower-quality descriptive studies "
    "tend to report that platforms boost innovation and partnerships, but the two strongest statistical "
    "studies found that platform-era biotech firms underperformed comparable companies financially, and the "
    "most common verdict overall is \"it depends.\" Almost no study tested when platforms help versus hurt, so "
    "the conditions for success remain unknown. The honest bottom line is that the current evidence is too "
    "thin, too varied in quality, and too inconsistent to support a confident claim that platform strategies "
    "improve biotechnology firm performance. The priority for future research is rigorous, design-justified, "
    "moderator-testing work that links platform adoption to financial as well as innovation outcomes."
)

# ---------------- REFERENCES ----------------
heading("6. References (included studies)", level=1)
para(
    "Author and year metadata were missing or unreliable for most source records; included studies are "
    "therefore listed by stable record ID and DOI where available, with MMAT score.",
    italic=True, size=9,
)
for e in included:
    qd = qmap.get(e["id"], {})
    doi = e.get("doi")
    yr = e.get("year") or "n.d."
    ref = f"{e['id']}"
    if doi:
        ref += f". https://doi.org/{doi}"
    ref += f" (MMAT {qd.get('total',0)}; method: {e.get('method') or 'not stated'}; year as recorded: {yr})."
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    for r in p.runs:
        r.font.size = Pt(9)

doc.save(str(OUT))
print(f"Saved {OUT}")
print(f"Included studies: {len(included)}")
